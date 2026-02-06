#!/usr/bin/env python3
"""
Weekly Vulnerability Comparison Report Generator
================================================
This script compares two vulnerability CSV exports and generates a comprehensive
Word document report showing changes by vendor, severity, and individual vulnerabilities.

Usage:
    python weekly_vuln_report.py <previous_week.csv> <current_week.csv> [output_name]

Example:
    python weekly_vuln_report.py "vulnerabilities-1-27.csv" "vulnerabilities-2-3.csv" "Weekly_Report_Feb3"

Requirements:
    pip install pandas python-docx
"""

import pandas as pd
import sys
import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import re
import json

# Chart generation
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import matplotlib.pyplot as plt

# =============================================================================
# CHART GENERATION FUNCTIONS
# =============================================================================

# Color schemes matching the report
COLORS = {
    'critical': '#7030A0',  # Purple
    'high': '#C00000',      # Dark Red
    'medium': '#ED7D31',    # Orange
    'low': '#70AD47',       # Green
    'cisco': '#1B4F72',     # Dark Blue
    'palo_alto': '#922B21', # Dark Red
    'other': '#7D7D7D',     # Gray
    'increase': '#C00000',  # Red
    'decrease': '#107C10',  # Green
    'neutral': '#2F5496',   # Blue
    'kev': '#B8860B',       # Dark Goldenrod for KEV
}

# =============================================================================
# KEV (KNOWN EXPLOITED VULNERABILITIES) FUNCTIONS
# =============================================================================

def load_kev_catalog(file_path="kev_catalog.json"):
    """Load CISA KEV catalog from local JSON file.

    Returns:
        dict: KEV catalog data or None if not found
    """
    # Try multiple filenames and locations
    filenames = [file_path, "known_exploited_vulnerabilities.json"]
    paths_to_try = []
    for fname in filenames:
        paths_to_try.append(fname)
        paths_to_try.append(os.path.join(os.path.dirname(__file__), fname))
        paths_to_try.append(os.path.join(os.getcwd(), fname))

    for path in paths_to_try:
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"Warning: Could not load KEV catalog from {path}: {e}")

    return None


def get_kev_cve_set(kev_data):
    """Extract set of CVE IDs from KEV catalog for fast lookup."""
    if not kev_data or 'vulnerabilities' not in kev_data:
        return set()
    return {v['cveID'] for v in kev_data['vulnerabilities'] if 'cveID' in v}


def get_kev_details(kev_data):
    """Get detailed KEV info as a dict keyed by CVE ID."""
    if not kev_data or 'vulnerabilities' not in kev_data:
        return {}
    return {v['cveID']: v for v in kev_data['vulnerabilities'] if 'cveID' in v}


def extract_cves_from_string(cve_string):
    """Extract individual CVE IDs from a comma-separated string."""
    if pd.isna(cve_string) or not cve_string:
        return []
    return re.findall(r'CVE-\d{4}-\d+', str(cve_string))


def analyze_kev_vulnerabilities(df_curr, kev_data):
    """Analyze current dataset against KEV catalog.

    Returns:
        dict with:
            - kev_cves: list of KEV CVEs found with asset counts
            - kev_assets: list of all assets with KEV vulnerabilities
            - total_kev_cves: count of unique KEV CVEs
            - total_kev_assets: count of unique assets with KEVs
    """
    if not kev_data:
        return None

    kev_set = get_kev_cve_set(kev_data)
    kev_details = get_kev_details(kev_data)

    if not kev_set:
        return None

    # Find all rows with KEV CVEs
    def has_kev(cve_string):
        cves = extract_cves_from_string(cve_string)
        return any(cve in kev_set for cve in cves)

    def get_matching_kevs(cve_string):
        cves = extract_cves_from_string(cve_string)
        return [cve for cve in cves if cve in kev_set]

    # Filter to rows with KEV vulnerabilities
    df_curr['has_kev'] = df_curr['definition.cve'].apply(has_kev)
    df_curr['kev_cves'] = df_curr['definition.cve'].apply(get_matching_kevs)

    kev_rows = df_curr[df_curr['has_kev']].copy()

    if len(kev_rows) == 0:
        return {
            'kev_cves': [],
            'kev_assets': [],
            'total_kev_cves': 0,
            'total_kev_assets': 0
        }

    # Build KEV CVE summary (CVE -> asset count, severity, vendor info)
    kev_cve_data = {}
    for _, row in kev_rows.iterrows():
        for cve in row['kev_cves']:
            if cve not in kev_cve_data:
                kev_info = kev_details.get(cve, {})
                kev_cve_data[cve] = {
                    'cve': cve,
                    'severity': row['severity'],
                    'vendor': row['vendor'],
                    'assets': set(),
                    'kev_vendor': kev_info.get('vendorProject', ''),
                    'kev_product': kev_info.get('product', ''),
                    'kev_name': kev_info.get('vulnerabilityName', '')[:60],
                    'date_added': kev_info.get('dateAdded', ''),
                    'due_date': kev_info.get('dueDate', '')
                }
            kev_cve_data[cve]['assets'].add(row['asset.id'])

    # Convert to list sorted by asset count
    kev_cves = []
    for cve, data in sorted(kev_cve_data.items(), key=lambda x: len(x[1]['assets']), reverse=True):
        kev_cves.append({
            'cve': data['cve'],
            'severity': data['severity'],
            'vendor': data['vendor'],
            'asset_count': len(data['assets']),
            'kev_name': data['kev_name'],
            'date_added': data['date_added'],
            'due_date': data['due_date']
        })

    # Build asset list with KEV details
    kev_assets = []
    asset_kev_map = {}  # asset_id -> list of KEV CVEs

    for _, row in kev_rows.iterrows():
        asset_id = row['asset.id']
        hostname = row['asset.host_name']

        if asset_id not in asset_kev_map:
            asset_kev_map[asset_id] = {
                'hostname': hostname if pd.notna(hostname) else row['asset.name'],
                'vendor': row['vendor'],
                'kev_cves': set(),
                'severities': set()
            }

        for cve in row['kev_cves']:
            asset_kev_map[asset_id]['kev_cves'].add(cve)
            asset_kev_map[asset_id]['severities'].add(row['severity'])

    # Build list with calculated fields first
    sev_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3}

    for asset_id, data in asset_kev_map.items():
        # Determine highest severity
        highest_sev = min(data['severities'], key=lambda s: sev_order.get(s, 99))
        hostname = str(data['hostname'])[:40]

        kev_assets.append({
            'hostname': hostname,
            'vendor': data['vendor'],
            'kev_count': len(data['kev_cves']),
            'kev_cves': ', '.join(sorted(data['kev_cves'])[:3]),  # First 3 CVEs
            'highest_severity': highest_sev,
            'sev_sort': sev_order.get(highest_sev, 99)  # For sorting
        })

    # Sort by: KEV count (desc), severity (asc), hostname (asc)
    kev_assets.sort(key=lambda x: (-x['kev_count'], x['sev_sort'], x['hostname'].lower()))

    return {
        'kev_cves': kev_cves,
        'kev_assets': kev_assets,
        'total_kev_cves': len(kev_cves),
        'total_kev_assets': len(kev_assets)
    }


# =============================================================================
# CVE CHANGE ANALYSIS FUNCTIONS
# =============================================================================

def analyze_cve_device_changes(df_prev, df_curr, cve_id):
    """Analyze which devices changed for a specific CVE."""
    # Filter to rows containing this CVE
    prev_cve = df_prev[df_prev['definition.cve'].str.contains(cve_id, na=False, regex=False)]
    curr_cve = df_curr[df_curr['definition.cve'].str.contains(cve_id, na=False, regex=False)]

    # Get unique hostnames
    prev_hosts = set(prev_cve['asset.host_name'].dropna().unique())
    curr_hosts = set(curr_cve['asset.host_name'].dropna().unique())

    # Calculate changes
    new_hosts = curr_hosts - prev_hosts
    removed_hosts = prev_hosts - curr_hosts

    return {
        'cve': cve_id,
        'prev_count': len(prev_hosts),
        'curr_count': len(curr_hosts),
        'new_hosts': sorted(new_hosts),
        'removed_hosts': sorted(removed_hosts),
        'new_count': len(new_hosts),
        'removed_count': len(removed_hosts)
    }


def analyze_top_cve_changes(df_prev, df_curr, vendor=None, top_n=5):
    """Get the CVEs with the biggest changes for a vendor."""
    # Filter by vendor if specified
    if vendor:
        df_p = df_prev[df_prev['vendor'] == vendor].copy()
        df_c = df_curr[df_curr['vendor'] == vendor].copy()
    else:
        df_p = df_prev.copy()
        df_c = df_curr.copy()

    # Count unique devices per CVE
    prev_counts = df_p.groupby(['definition.cve', 'definition.name', 'severity']).agg({
        'asset.host_name': 'nunique'
    }).reset_index()
    prev_counts.columns = ['cve', 'name', 'severity', 'prev_devices']

    curr_counts = df_c.groupby(['definition.cve', 'definition.name', 'severity']).agg({
        'asset.host_name': 'nunique'
    }).reset_index()
    curr_counts.columns = ['cve', 'name', 'severity', 'curr_devices']

    # Merge and calculate change
    merged = prev_counts.merge(curr_counts, on=['cve', 'name', 'severity'], how='outer').fillna(0)
    merged['change'] = merged['curr_devices'] - merged['prev_devices']

    # Get top increases and decreases
    increases = merged[merged['change'] > 0].nlargest(top_n, 'change')
    decreases = merged[merged['change'] < 0].nsmallest(top_n, 'change')

    return increases, decreases


def create_severity_comparison_chart(results, temp_dir):
    """Create a bar chart comparing previous vs current by severity."""
    severities = ['Critical', 'High', 'Medium', 'Low']
    prev_values = [results['by_severity'].get(s, {}).get('prev', 0) for s in severities]
    curr_values = [results['by_severity'].get(s, {}).get('curr', 0) for s in severities]

    fig, ax = plt.subplots(figsize=(8, 4))

    x = range(len(severities))
    width = 0.35

    bars1 = ax.bar([i - width/2 for i in x], prev_values, width, label='Previous', color='#5B9BD5', edgecolor='white')
    bars2 = ax.bar([i + width/2 for i in x], curr_values, width, label='Current', color='#2F5496', edgecolor='white')

    ax.set_ylabel('Vulnerability Count', fontsize=11, fontweight='bold')
    ax.set_title('Vulnerability Comparison by Severity', fontsize=14, fontweight='bold', pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(severities, fontsize=10)
    ax.legend(loc='upper right', framealpha=0.9)

    # Add value labels on bars
    for bar in bars1:
        height = bar.get_height()
        if height > 0:
            ax.annotate(f'{int(height):,}', xy=(bar.get_x() + bar.get_width()/2, height),
                       xytext=(0, 3), textcoords="offset points", ha='center', va='bottom', fontsize=8)
    for bar in bars2:
        height = bar.get_height()
        if height > 0:
            ax.annotate(f'{int(height):,}', xy=(bar.get_x() + bar.get_width()/2, height),
                       xytext=(0, 3), textcoords="offset points", ha='center', va='bottom', fontsize=8)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.yaxis.grid(True, linestyle='--', alpha=0.7)
    ax.set_axisbelow(True)

    plt.tight_layout()
    chart_path = os.path.join(temp_dir, 'severity_chart.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return chart_path


def create_vendor_pie_chart(results, temp_dir):
    """Create a pie chart showing vulnerability distribution by vendor."""
    vendors = []
    values = []
    colors = []

    color_map = {
        'Cisco': '#1B4F72',
        'Palo Alto': '#922B21',
        'Other/General': '#7D7D7D'
    }

    for vendor, data in sorted(results['by_vendor'].items(), key=lambda x: x[1]['curr'], reverse=True):
        if data['curr'] > 0:
            vendors.append(vendor)
            values.append(data['curr'])
            colors.append(color_map.get(vendor, '#5B9BD5'))

    fig, ax = plt.subplots(figsize=(6, 4))

    wedges, texts, autotexts = ax.pie(values, labels=vendors, autopct='%1.1f%%',
                                       colors=colors, startangle=90,
                                       explode=[0.02] * len(vendors),
                                       shadow=False,
                                       textprops={'fontsize': 10})

    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontweight('bold')

    ax.set_title('Current Vulnerabilities by Vendor', fontsize=14, fontweight='bold', pad=15)

    plt.tight_layout()
    chart_path = os.path.join(temp_dir, 'vendor_pie_chart.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return chart_path


def create_change_summary_chart(results, temp_dir):
    """Create a horizontal bar chart showing changes by severity."""
    severities = ['Critical', 'High', 'Medium', 'Low']
    changes = [results['by_severity'].get(s, {}).get('change', 0) for s in severities]

    colors = ['#C00000' if c > 0 else '#107C10' if c < 0 else '#7D7D7D' for c in changes]

    fig, ax = plt.subplots(figsize=(7, 3))

    y_pos = range(len(severities))
    bars = ax.barh(y_pos, changes, color=colors, edgecolor='white', height=0.6)

    ax.set_yticks(y_pos)
    ax.set_yticklabels(severities, fontsize=10)
    ax.set_xlabel('Change in Vulnerability Count', fontsize=11, fontweight='bold')
    ax.set_title('Week-over-Week Change by Severity', fontsize=14, fontweight='bold', pad=15)

    # Add value labels
    for i, (bar, val) in enumerate(zip(bars, changes)):
        if val != 0:
            label = f'+{val:,}' if val > 0 else f'{val:,}'
            x_pos = val + (50 if val > 0 else -50)
            ax.annotate(label, xy=(val, bar.get_y() + bar.get_height()/2),
                       xytext=(5 if val > 0 else -5, 0), textcoords="offset points",
                       ha='left' if val > 0 else 'right', va='center', fontsize=9, fontweight='bold')

    ax.axvline(x=0, color='black', linewidth=0.8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.xaxis.grid(True, linestyle='--', alpha=0.7)
    ax.set_axisbelow(True)

    plt.tight_layout()
    chart_path = os.path.join(temp_dir, 'change_chart.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return chart_path


def add_table_of_contents(doc):
    """Add a manual Table of Contents to the document."""
    # Define TOC entries with section names
    toc_entries = [
        ("Executive Summary", 2),
        ("    Overall Summary", 2),
        ("    Breakdown by Severity", 2),
        ("    Breakdown by Vendor", 2),
        ("Detailed Breakdown: Vendor x Severity", 3),
        ("Top Vulnerability Changes by Vendor", 4),
        ("Current Top Vulnerabilities by Vendor", 5),
        ("Asset Risk Analysis", 6),
        ("    Most Vulnerable Assets", 6),
        ("    Assets with Critical Vulnerabilities", 6),
        ("KEV Analysis (if available)", 7),
        ("Definition & Asset Changes", 8),
        ("CVE Change Analysis", 9),
    ]

    # Create TOC table for clean formatting
    table = doc.add_table(rows=len(toc_entries), cols=2)

    for i, (section, page) in enumerate(toc_entries):
        row = table.rows[i]
        # Section name
        cell1 = row.cells[0]
        run1 = cell1.paragraphs[0].add_run(section)
        if not section.startswith("    "):
            run1.bold = True
        run1.font.size = Pt(11)

        # Page number
        cell2 = row.cells[1]
        run2 = cell2.paragraphs[0].add_run(str(page))
        run2.font.size = Pt(11)
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

# =============================================================================
# HOSTNAME AND VENDOR NORMALIZATION (Matching HexTracker methodology)
# =============================================================================

def normalize_hostname(hostname):
    """Normalize hostname for consistent processing.

    Matches HexTracker's normalizeHostname() function:
    - For IP addresses, return the full IP
    - For domain names, remove everything after first period

    Examples:
        nwan10.mmplp.net -> nwan10
        192.168.1.1 -> 192.168.1.1
    """
    if not hostname or pd.isna(hostname):
        return ""

    clean_hostname = str(hostname).strip()

    # Check if hostname is a valid IP address
    ip_pattern = r'^(\d{1,3}\.){3}\d{1,3}$'
    if re.match(ip_pattern, clean_hostname):
        octets = clean_hostname.split('.')
        try:
            if all(0 <= int(octet) <= 255 for octet in octets):
                return clean_hostname.lower()
        except ValueError:
            pass

    # For domain names, remove everything after first period
    return clean_hostname.split('.')[0].lower()


def categorize_device_type(hostname, family=None):
    """Categorize device by hostname pattern with family fallback.

    Matches HexTracker's device-naming-patterns.json configuration.
    Patterns are checked in precedence order.

    Cisco patterns:
    - nswan, swan - Switches
    - nrwan, rtr - Routers
    - fw, asa - Firewalls
    - sw - Generic switch
    - newan - Edge devices

    Palo Alto patterns:
    - nfpan, dnfpan - Network firewalls
    - nfscada - SCADA firewalls
    - nfdevpan, nfextpan, nfintpan, nfcampuspan - Specific firewall types
    - papan - PA firewalls
    - extpan, intpan, campuspan, devpan - Generic firewall types

    Other patterns:
    - nwwbr, wwbr - Wireless bridges
    """
    # Normalize hostname first (remove domain suffix)
    h = normalize_hostname(hostname)
    f = str(family).lower() if family and pd.notna(family) else ""

    # Step 1: Hostname patterns (from HexTracker device-naming-patterns.json)
    # Patterns checked in precedence order
    if h:
        # Cisco patterns (precedence 1-5, 17-21)
        if 'nswan' in h:
            return 'Cisco'
        if 'swan' in h:
            return 'Cisco'
        if 'nrwan' in h:
            return 'Cisco'
        if 'rtr' in h:
            return 'Cisco'
        if re.search(r'ns\d+com01', h):
            return 'Cisco'
        if 'fw' in h and 'nf' not in h:  # fw but not nfpan patterns
            return 'Cisco'
        if 'asa' in h:
            return 'Cisco'
        if 'sw' in h and 'wan' not in h:  # sw but not swan
            return 'Cisco'
        if 'newan' in h:
            return 'Cisco'

        # Palo Alto patterns (precedence 6-16) - check specific patterns first
        if 'nfpan' in h:
            return 'Palo Alto'
        if 'dnfpan' in h:
            return 'Palo Alto'
        if 'nfscada' in h:
            return 'Palo Alto'
        if 'nfdevpan' in h:
            return 'Palo Alto'
        if 'papan' in h:
            return 'Palo Alto'
        if 'nfextpan' in h:
            return 'Palo Alto'
        if 'nfintpan' in h:
            return 'Palo Alto'
        if 'nfcampuspan' in h:
            return 'Palo Alto'
        if 'extpan' in h:
            return 'Palo Alto'
        if 'intpan' in h:
            return 'Palo Alto'
        if 'campuspan' in h:
            return 'Palo Alto'
        if 'devpan' in h:
            return 'Palo Alto'

        # Other patterns (precedence 22-23) - wireless bridges
        if 'nwwbr' in h or 'wwbr' in h:
            return 'Other/General'

    # Step 2: Family patterns (HexTracker familyVendorPatterns fallback)
    if f:
        if 'cisco' in f:
            return 'Cisco'
        if re.search(r'palo\s*alto', f):
            return 'Palo Alto'

    return 'Other/General'


def normalize_vendor(family):
    """Normalize vendor names from definition.family column.
    DEPRECATED: Use categorize_device_type() instead for device-based categorization.
    """
    family_upper = str(family).upper()
    if 'CISCO' in family_upper:
        return 'Cisco'
    elif 'PALO' in family_upper:
        return 'Palo Alto'
    else:
        return 'Other/General'

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<w:r><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
        f'<w:t>{text}</w:t></w:r></w:hyperlink>'
    )
    paragraph._p.append(hyperlink)
    return hyperlink

def get_cve_url(cve_text):
    """Generate CVE.org URL from CVE identifier."""
    # Extract CVE ID pattern (CVE-YYYY-NNNNN)
    match = re.search(r'(CVE-\d{4}-\d+)', str(cve_text))
    if match:
        return f"https://www.cve.org/CVERecord?id={match.group(1)}"
    return None

def add_styled_table(doc, headers, rows, header_color="2F5496", hyperlink_col=None, font_size=None):
    """Add a styled table to the document.

    Args:
        hyperlink_col: Column index to make into CVE hyperlinks (0-based), or None for no hyperlinks
        font_size: Optional font size in points for table content (default uses Word default)
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        run = header_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        if font_size:
            run.font.size = Pt(font_size)
        set_cell_shading(header_cells[i], header_color)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            if hyperlink_col is not None and i == hyperlink_col:
                # Make this cell a hyperlink to CVE.org
                cve_url = get_cve_url(value)
                if cve_url:
                    # Clear default paragraph and add hyperlink
                    paragraph = row_cells[i].paragraphs[0]
                    add_hyperlink(paragraph, str(value), cve_url)
                else:
                    row_cells[i].text = str(value)
                    if font_size:
                        row_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
            else:
                row_cells[i].text = str(value)
                if font_size and row_cells[i].paragraphs[0].runs:
                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)

    return table

def add_styled_table_with_change_colors(doc, headers, rows, change_colors, header_color="2F5496", hyperlink_col=None):
    """Add a styled table with colored Change column (last column).

    Args:
        change_colors: List of hex color codes for each row's change value
        hyperlink_col: Column index to make into CVE hyperlinks (0-based), or None for no hyperlinks
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            # Handle hyperlink column
            if hyperlink_col is not None and i == hyperlink_col:
                cve_url = get_cve_url(value)
                if cve_url:
                    paragraph = row_cells[i].paragraphs[0]
                    add_hyperlink(paragraph, str(value), cve_url)
                else:
                    row_cells[i].text = str(value)
            else:
                row_cells[i].text = str(value)

            # Color and bold the last column (Change)
            if i == len(row_data) - 1 and row_idx < len(change_colors):
                run = row_cells[i].paragraphs[0].runs[0]
                run.bold = True
                color_hex = change_colors[row_idx]
                run.font.color.rgb = RGBColor(
                    int(color_hex[0:2], 16),
                    int(color_hex[2:4], 16),
                    int(color_hex[4:6], 16)
                )

    return table

def analyze_data(df_prev, df_curr):
    """Perform comprehensive analysis on the two datasets."""
    results = {}

    # Add vendor column using HexTracker-style categorization
    # Uses hostname patterns first, then falls back to definition.family
    df_prev['vendor'] = df_prev.apply(
        lambda row: categorize_device_type(row['asset.host_name'], row.get('definition.family')),
        axis=1
    )
    df_curr['vendor'] = df_curr.apply(
        lambda row: categorize_device_type(row['asset.host_name'], row.get('definition.family')),
        axis=1
    )

    # Overall summary
    results['overall'] = {
        'prev_total': len(df_prev),
        'curr_total': len(df_curr),
        'prev_assets': df_prev['asset.id'].nunique(),
        'curr_assets': df_curr['asset.id'].nunique(),
        'prev_definitions': df_prev['definition.id'].nunique(),
        'curr_definitions': df_curr['definition.id'].nunique(),
    }

    # By severity
    severities = ['Critical', 'High', 'Medium', 'Low']
    results['by_severity'] = {}
    for sev in severities:
        prev_count = len(df_prev[df_prev['severity'] == sev])
        curr_count = len(df_curr[df_curr['severity'] == sev])
        results['by_severity'][sev] = {'prev': prev_count, 'curr': curr_count, 'change': curr_count - prev_count}

    # By vendor
    vendors = df_curr['vendor'].unique().tolist()
    results['by_vendor'] = {}
    for vendor in vendors:
        prev_count = len(df_prev[df_prev['vendor'] == vendor])
        curr_count = len(df_curr[df_curr['vendor'] == vendor])
        results['by_vendor'][vendor] = {'prev': prev_count, 'curr': curr_count, 'change': curr_count - prev_count}

    # Vendor x Severity breakdown
    results['vendor_severity'] = {}
    for vendor in vendors:
        results['vendor_severity'][vendor] = {}
        for sev in severities:
            prev_count = len(df_prev[(df_prev['vendor'] == vendor) & (df_prev['severity'] == sev)])
            curr_count = len(df_curr[(df_curr['vendor'] == vendor) & (df_curr['severity'] == sev)])
            if prev_count > 0 or curr_count > 0:
                results['vendor_severity'][vendor][sev] = {'prev': prev_count, 'curr': curr_count, 'change': curr_count - prev_count}

    # New and removed definitions
    new_def_ids = set(df_curr['definition.id']) - set(df_prev['definition.id'])
    removed_def_ids = set(df_prev['definition.id']) - set(df_curr['definition.id'])

    results['new_definitions'] = []
    if new_def_ids:
        new_defs = df_curr[df_curr['definition.id'].isin(new_def_ids)][['definition.name', 'definition.cve', 'severity', 'vendor']].drop_duplicates()
        for _, row in new_defs.iterrows():
            results['new_definitions'].append({
                'name': row['definition.name'][:70],
                'cve': str(row['definition.cve'])[:50],
                'severity': row['severity'],
                'vendor': row['vendor']
            })

    results['removed_definitions'] = []
    if removed_def_ids:
        removed_defs = df_prev[df_prev['definition.id'].isin(removed_def_ids)][['definition.name', 'definition.cve', 'severity', 'vendor']].drop_duplicates()
        for _, row in removed_defs.iterrows():
            results['removed_definitions'].append({
                'name': row['definition.name'][:70],
                'cve': str(row['definition.cve'])[:50],
                'severity': row['severity'],
                'vendor': row['vendor']
            })

    # Asset changes
    new_assets = set(df_curr['asset.id']) - set(df_prev['asset.id'])
    removed_assets = set(df_prev['asset.id']) - set(df_curr['asset.id'])

    results['new_assets'] = []
    if new_assets:
        new_asset_data = df_curr[df_curr['asset.id'].isin(new_assets)][['asset.host_name', 'asset.name']].drop_duplicates()
        # Use host_name if available, fall back to asset.name
        for _, row in new_asset_data.iterrows():
            hostname = row['asset.host_name']
            asset_name = row['asset.name']
            if pd.notna(hostname) and str(hostname).strip() and str(hostname).lower() != 'nan':
                results['new_assets'].append(str(hostname))
            elif pd.notna(asset_name) and str(asset_name).strip() and str(asset_name).lower() != 'nan':
                results['new_assets'].append(f"{asset_name} (no hostname)")
        results['new_assets'] = results['new_assets'][:20]

    results['removed_assets'] = []
    if removed_assets:
        removed_asset_data = df_prev[df_prev['asset.id'].isin(removed_assets)][['asset.host_name', 'asset.name']].drop_duplicates()
        # Use host_name if available, fall back to asset.name
        for _, row in removed_asset_data.iterrows():
            hostname = row['asset.host_name']
            asset_name = row['asset.name']
            if pd.notna(hostname) and str(hostname).strip() and str(hostname).lower() != 'nan':
                results['removed_assets'].append(str(hostname))
            elif pd.notna(asset_name) and str(asset_name).strip() and str(asset_name).lower() != 'nan':
                results['removed_assets'].append(f"{asset_name} (no hostname)")
        results['removed_assets'] = results['removed_assets'][:20]

    results['new_asset_count'] = len(new_assets)
    results['removed_asset_count'] = len(removed_assets)

    # Top vulnerability changes by vendor (top 5 increases and decreases per vendor)
    count_prev = df_prev.groupby(['definition.id', 'definition.name', 'definition.cve', 'severity', 'vendor']).size().reset_index(name='count_prev')
    count_curr = df_curr.groupby(['definition.id', 'definition.name', 'definition.cve', 'severity', 'vendor']).size().reset_index(name='count_curr')
    comparison = count_prev.merge(count_curr, on=['definition.id', 'definition.name', 'definition.cve', 'severity', 'vendor'], how='outer').fillna(0)
    comparison['change'] = comparison['count_curr'] - comparison['count_prev']

    # Get vendor-specific top changes (Cisco and Palo Alto only)
    results['vendor_changes'] = {}
    for vendor in ['Cisco', 'Palo Alto']:
        vendor_data = comparison[comparison['vendor'] == vendor]

        increases = []
        for _, row in vendor_data.nlargest(5, 'change').iterrows():
            if row['change'] > 0:
                cve = str(row['definition.cve']).split(',')[0].strip()[:30]
                increases.append({
                    'name': row['definition.name'][:60],
                    'cve': cve,
                    'severity': row['severity'],
                    'prev': int(row['count_prev']),
                    'curr': int(row['count_curr']),
                    'change': int(row['change'])
                })

        decreases = []
        for _, row in vendor_data.nsmallest(5, 'change').iterrows():
            if row['change'] < 0:
                cve = str(row['definition.cve']).split(',')[0].strip()[:30]
                decreases.append({
                    'name': row['definition.name'][:60],
                    'cve': cve,
                    'severity': row['severity'],
                    'prev': int(row['count_prev']),
                    'curr': int(row['count_curr']),
                    'change': int(row['change'])
                })

        results['vendor_changes'][vendor] = {
            'increases': increases,
            'decreases': decreases
        }

    # Current top vulnerabilities by vendor and severity (from current dataset only)
    results['current_top_vulns'] = {}
    for vendor in ['Cisco', 'Palo Alto']:
        results['current_top_vulns'][vendor] = {}
        vendor_curr = df_curr[df_curr['vendor'] == vendor]

        for severity in ['Critical', 'High', 'Medium']:
            sev_data = vendor_curr[vendor_curr['severity'] == severity]
            if len(sev_data) == 0:
                results['current_top_vulns'][vendor][severity] = []
                continue

            # Group by vulnerability and count unique assets
            top_vulns = sev_data.groupby(['definition.id', 'definition.name', 'definition.cve']).agg({
                'asset.id': 'nunique'
            }).reset_index()
            top_vulns.columns = ['definition.id', 'definition.name', 'definition.cve', 'asset_count']
            top_vulns = top_vulns.nlargest(5, 'asset_count')

            vuln_list = []
            for _, row in top_vulns.iterrows():
                cve = str(row['definition.cve']).split(',')[0].strip()[:30]
                # Get a truncated description from the vulnerability name
                vuln_name = str(row['definition.name'])
                # Clean up the name - remove the cisco-sa reference in parentheses for brevity
                if '(' in vuln_name:
                    description = vuln_name.split('(')[0].strip()[:60]
                else:
                    description = vuln_name[:60]
                vuln_list.append({
                    'cve': cve,
                    'name': vuln_name,
                    'description': description,
                    'asset_count': int(row['asset_count'])
                })
            results['current_top_vulns'][vendor][severity] = vuln_list

    # Top 5 Most Vulnerable Assets by CVE-Weighted VPR (per vendor)
    # This matches Hextrackr methodology: VPR × CVE count per definition
    def count_cves(cve_str):
        """Count individual CVEs in a comma-separated string."""
        if pd.isna(cve_str):
            return 1
        cves = [c for c in str(cve_str).split(',') if 'CVE' in c]
        return max(len(cves), 1)

    results['most_vulnerable_assets'] = {}
    for vendor in ['Cisco', 'Palo Alto']:
        vendor_curr = df_curr[df_curr['vendor'] == vendor].copy()

        # Calculate CVE count and weighted VPR for each row
        vendor_curr['cve_count'] = vendor_curr['definition.cve'].apply(count_cves)
        vendor_curr['weighted_vpr'] = vendor_curr['definition.vpr.score'] * vendor_curr['cve_count']

        # Group by asset and sum
        asset_vpr = vendor_curr.groupby(['asset.id', 'asset.host_name', 'asset.name']).agg({
            'weighted_vpr': 'sum',
            'cve_count': 'sum'
        }).reset_index()
        asset_vpr.columns = ['asset_id', 'hostname', 'asset_name', 'vpr_sum', 'cve_count']
        asset_vpr = asset_vpr.nlargest(5, 'vpr_sum')

        asset_list = []
        for _, row in asset_vpr.iterrows():
            hostname = row['hostname']
            asset_name = row['asset_name']
            if pd.notna(hostname) and str(hostname).strip() and str(hostname).lower() != 'nan':
                display_name = str(hostname)
            elif pd.notna(asset_name) and str(asset_name).strip():
                display_name = f"{asset_name} (no hostname)"
            else:
                display_name = "Unknown"

            asset_list.append({
                'name': display_name,
                'vpr_sum': round(row['vpr_sum'], 1),
                'cve_count': int(row['cve_count'])
            })
        results['most_vulnerable_assets'][vendor] = asset_list

    # All Assets with Critical Vulnerabilities
    critical = df_curr[df_curr['severity'] == 'Critical']
    critical_assets = critical.groupby(['asset.host_name', 'asset.name', 'vendor']).agg({
        'definition.id': 'nunique',
        'definition.vpr.score': 'sum',
        'definition.cve': lambda x: ', '.join(x.unique()[:3])  # First 3 CVEs
    }).reset_index()
    critical_assets.columns = ['hostname', 'asset_name', 'vendor', 'crit_count', 'vpr_sum', 'cves']

    results['critical_assets'] = []
    for _, row in critical_assets.iterrows():
        hostname = row['hostname']
        asset_name = row['asset_name']
        if pd.notna(hostname) and str(hostname).strip() and str(hostname).lower() != 'nan':
            display_name = str(hostname)
        elif pd.notna(asset_name) and str(asset_name).strip():
            display_name = f"{asset_name} (no hostname)"
        else:
            display_name = "Unknown"

        # Only include Cisco and Palo Alto
        if row['vendor'] in ['Cisco', 'Palo Alto']:
            results['critical_assets'].append({
                'name': display_name,
                'vendor': row['vendor'],
                'crit_count': int(row['crit_count']),
                'vpr_sum': round(row['vpr_sum'], 1),
                'cves': str(row['cves']).split(',')[0].strip()[:30]  # First CVE
            })

    # Sort by VPR sum descending
    results['critical_assets'] = sorted(results['critical_assets'], key=lambda x: x['vpr_sum'], reverse=True)

    return results, df_prev, df_curr

def generate_report(results, output_path, prev_name, curr_name):
    """Generate the Word document report with professional formatting."""
    doc = Document()

    # Create temporary directory for charts
    temp_dir = tempfile.mkdtemp()

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    title = doc.add_heading('Netops Weekly Vulnerability Comparison Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with dates
    subtitle = doc.add_paragraph(f'Comparing: {prev_name} vs {curr_name}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    toc_heading = doc.add_heading('Table of Contents', level=1)
    add_table_of_contents(doc)

    doc.add_page_break()

    # =========================================================================
    # EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading('Executive Summary', level=1)

    overall = results['overall']
    total_change = overall['curr_total'] - overall['prev_total']
    change_word = "increase" if total_change > 0 else "decrease" if total_change < 0 else "no change"

    summary = doc.add_paragraph()
    summary.add_run(f"This report compares vulnerability scan data between two periods. ")
    summary.add_run(f"The current dataset shows a net {change_word} of {abs(total_change):,} vulnerability instances. ")
    summary.add_run(f"Asset inventory changed by {results['new_asset_count'] - results['removed_asset_count']:+d} ({results['new_asset_count']} added, {results['removed_asset_count']} removed).")

    # -------------------------------------------------------------------------
    # Executive Summary Charts
    # -------------------------------------------------------------------------
    doc.add_paragraph()

    # Severity Comparison Chart
    try:
        severity_chart_path = create_severity_comparison_chart(results, temp_dir)
        doc.add_picture(severity_chart_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Chart could not be generated: {e}]")

    doc.add_paragraph()

    # Change Summary Chart (horizontal bar)
    try:
        change_chart_path = create_change_summary_chart(results, temp_dir)
        doc.add_picture(change_chart_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Chart could not be generated: {e}]")

    doc.add_page_break()

    # =========================================================================
    # SUMMARY TABLES
    # =========================================================================

    # Overall Summary Table
    doc.add_heading('Overall Summary', level=2)

    headers = ['Metric', 'Previous', 'Current', 'Change']
    # Calculate changes
    vuln_change = overall['curr_total'] - overall['prev_total']
    asset_change = overall['curr_assets'] - overall['prev_assets']
    def_change = overall['curr_definitions'] - overall['prev_definitions']

    rows = [
        ['Total Vulnerabilities', f"{overall['prev_total']:,}", f"{overall['curr_total']:,}", f"{vuln_change:+,}"],
        ['Unique Assets', f"{overall['prev_assets']:,}", f"{overall['curr_assets']:,}", f"{asset_change:+,}"],
        ['Unique Definitions', f"{overall['prev_definitions']:,}", f"{overall['curr_definitions']:,}", f"{def_change:+,}"],
    ]
    # For vulnerabilities: more = bad (red), fewer = good (green)
    # For assets/definitions: neutral coloring (more assets isn't necessarily bad)
    change_colors = []
    for change in [vuln_change, asset_change, def_change]:
        if change > 0:
            change_colors.append("C00000")  # Red
        elif change < 0:
            change_colors.append("107C10")  # Green
        else:
            change_colors.append("000000")  # Black
    add_styled_table_with_change_colors(doc, headers, rows, change_colors)

    doc.add_paragraph()

    # By Severity
    doc.add_heading('Breakdown by Severity', level=2)

    headers = ['Severity', 'Previous', 'Current', 'Change']
    rows = []
    change_colors = []
    for sev in ['Critical', 'High', 'Medium', 'Low']:
        data = results['by_severity'].get(sev, {'prev': 0, 'curr': 0, 'change': 0})
        change_val = data['change']
        rows.append([sev, f"{data['prev']:,}", f"{data['curr']:,}", f"{change_val:+,}"])
        # Color: red for increase (bad), green for decrease (good), black for no change
        if change_val > 0:
            change_colors.append("C00000")  # Red
        elif change_val < 0:
            change_colors.append("107C10")  # Green
        else:
            change_colors.append("000000")  # Black
    add_styled_table_with_change_colors(doc, headers, rows, change_colors)

    doc.add_paragraph()

    # By Vendor
    doc.add_heading('Breakdown by Vendor', level=2)

    headers = ['Vendor', 'Previous', 'Current', 'Change']
    rows = []
    change_colors = []
    for vendor, data in sorted(results['by_vendor'].items(), key=lambda x: x[1]['curr'], reverse=True):
        change_val = data['change']
        rows.append([vendor, f"{data['prev']:,}", f"{data['curr']:,}", f"{change_val:+,}"])
        # Color: red for increase (bad), green for decrease (good), black for no change
        if change_val > 0:
            change_colors.append("C00000")  # Red
        elif change_val < 0:
            change_colors.append("107C10")  # Green
        else:
            change_colors.append("000000")  # Black
    add_styled_table_with_change_colors(doc, headers, rows, change_colors)

    doc.add_paragraph()

    # Vendor Distribution Pie Chart (placed after Breakdown by Vendor table)
    try:
        vendor_chart_path = create_vendor_pie_chart(results, temp_dir)
        doc.add_picture(vendor_chart_path, width=Inches(3.75))  # Reduced ~25% to fit on same page
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Chart could not be generated: {e}]")

    doc.add_page_break()

    # =========================================================================
    # DETAILED BREAKDOWNS
    # =========================================================================
    doc.add_heading('Detailed Breakdown: Vendor x Severity', level=2)

    # Show only Cisco and Palo Alto with all 4 severity levels
    for vendor in ['Cisco', 'Palo Alto']:
        sev_data = results['vendor_severity'].get(vendor, {})
        doc.add_heading(vendor, level=3)
        headers = ['Severity', 'Previous', 'Current', 'Change']
        rows = []
        change_colors = []  # Track colors for change column
        for sev in ['Critical', 'High', 'Medium', 'Low']:
            if sev in sev_data:
                data = sev_data[sev]
                change_val = data['change']
                rows.append([sev, f"{data['prev']:,}", f"{data['curr']:,}", f"{change_val:+,}"])
            else:
                # Show 0s for missing severity levels
                change_val = 0
                rows.append([sev, "0", "0", "0"])
            # Determine color: green for positive, red for negative, black for zero
            if change_val > 0:
                change_colors.append("C00000")  # Red (more vulns = bad)
            elif change_val < 0:
                change_colors.append("107C10")  # Green (fewer vulns = good)
            else:
                change_colors.append("000000")  # Black (no change)
        add_styled_table_with_change_colors(doc, headers, rows, change_colors, header_color="4472C4")
        doc.add_paragraph()

    doc.add_page_break()

    # =========================================================================
    # TOP VULNERABILITY CHANGES
    # =========================================================================
    doc.add_heading('Top Vulnerability Changes by Vendor', level=2)

    vendor_colors = {
        'Cisco': {'increase': 'C65911', 'decrease': '548235'},
        'Palo Alto': {'increase': 'C65911', 'decrease': '548235'}
    }

    for vendor in ['Cisco', 'Palo Alto']:
        vendor_data = results['vendor_changes'].get(vendor, {'increases': [], 'decreases': []})

        doc.add_heading(f'{vendor}', level=3)

        # Increases
        if vendor_data['increases']:
            doc.add_paragraph().add_run('Top Increases:').bold = True
            headers = ['CVE', 'Severity', 'Prev', 'Curr', 'Change']
            rows = []
            change_colors = []
            for item in vendor_data['increases']:
                rows.append([item['cve'], item['severity'],
                            str(item['prev']), str(item['curr']), f"+{item['change']}"])
                change_colors.append("C00000")  # Red for increases (bad)
            add_styled_table_with_change_colors(doc, headers, rows, change_colors,
                                                header_color=vendor_colors[vendor]['increase'], hyperlink_col=0)
        else:
            doc.add_paragraph('No increases this period.', style='List Bullet')

        doc.add_paragraph()

        # Decreases
        if vendor_data['decreases']:
            doc.add_paragraph().add_run('Top Decreases:').bold = True
            headers = ['CVE', 'Severity', 'Prev', 'Curr', 'Change']
            rows = []
            change_colors = []
            for item in vendor_data['decreases']:
                rows.append([item['cve'], item['severity'],
                            str(item['prev']), str(item['curr']), str(item['change'])])
                change_colors.append("107C10")  # Green for decreases (good)
            add_styled_table_with_change_colors(doc, headers, rows, change_colors,
                                                header_color=vendor_colors[vendor]['decrease'], hyperlink_col=0)
        else:
            doc.add_paragraph('No decreases this period.', style='List Bullet')

        doc.add_paragraph()

    doc.add_page_break()

    # =========================================================================
    # CURRENT TOP VULNERABILITIES
    # =========================================================================
    doc.add_heading('Current Top Vulnerabilities by Vendor', level=2)

    severity_colors = {
        'Critical': '7030A0',  # Purple
        'High': 'C00000',      # Dark Red
        'Medium': 'ED7D31'     # Orange
    }

    for vendor in ['Cisco', 'Palo Alto']:
        doc.add_heading(f'{vendor}', level=3)
        vendor_vulns = results.get('current_top_vulns', {}).get(vendor, {})

        for severity in ['Critical', 'High', 'Medium']:
            sev_vulns = vendor_vulns.get(severity, [])

            para = doc.add_paragraph()
            para.add_run(f'{severity}: ').bold = True

            if sev_vulns:
                headers = ['CVE', 'Assets', 'Description']
                rows = [[v['cve'], str(v['asset_count']), v['description']] for v in sev_vulns]
                add_styled_table(doc, headers, rows, header_color=severity_colors[severity], hyperlink_col=0)
            else:
                doc.add_paragraph(f'No {severity.lower()} vulnerabilities found.', style='List Bullet')

            doc.add_paragraph()

    doc.add_page_break()

    # =========================================================================
    # ASSET RISK ANALYSIS
    # =========================================================================
    doc.add_heading('Most Vulnerable Assets (by VPR Score)', level=2)

    for vendor in ['Cisco', 'Palo Alto']:
        doc.add_heading(f'{vendor}', level=3)
        assets = results.get('most_vulnerable_assets', {}).get(vendor, [])

        if assets:
            headers = ['Asset', 'VPR Sum', 'CVE Count']
            rows = [[a['name'], str(a['vpr_sum']), str(a['cve_count'])] for a in assets]
            add_styled_table(doc, headers, rows, header_color="1F4E79")
        else:
            doc.add_paragraph('No assets found.', style='List Bullet')

        doc.add_paragraph()

    # Page break before Assets with Critical Vulnerabilities
    doc.add_page_break()

    # Assets with Critical Vulnerabilities
    doc.add_heading('Assets with Critical Vulnerabilities', level=2)

    critical_assets = results.get('critical_assets', [])
    if critical_assets:
        para = doc.add_paragraph()
        para.add_run(f'Total: {len(critical_assets)} assets with critical vulnerabilities').bold = True
        doc.add_paragraph()

        headers = ['Asset', 'Vendor', 'Critical CVE', 'VPR']
        rows = [[a['name'], a['vendor'], a['cves'], str(a['vpr_sum'])] for a in critical_assets]
        add_styled_table(doc, headers, rows, header_color="7030A0", hyperlink_col=2)
    else:
        doc.add_paragraph('No assets with critical vulnerabilities found.')

    doc.add_paragraph()

    # =========================================================================
    # KEV (KNOWN EXPLOITED VULNERABILITIES) ANALYSIS
    # =========================================================================
    kev_results = results.get('kev_analysis')
    if kev_results and kev_results.get('total_kev_cves', 0) > 0:
        doc.add_page_break()
        doc.add_heading('KEV (Known Exploited Vulnerabilities) Analysis', level=1)

        # Summary paragraph
        summary = doc.add_paragraph()
        summary.add_run(f"CISA Known Exploited Vulnerabilities found in environment: ").bold = True
        summary.add_run(f"{kev_results['total_kev_cves']} KEV CVEs affecting {kev_results['total_kev_assets']} assets. ")
        summary.add_run("These vulnerabilities have confirmed active exploitation and should be prioritized for remediation.")

        doc.add_paragraph()

        # KEV CVEs by Impact (summary table)
        doc.add_heading('KEV CVEs by Asset Impact', level=2)

        kev_cves = kev_results.get('kev_cves', [])
        if kev_cves:
            headers = ['CVE', 'Severity', 'Vendor', 'Assets', 'KEV Name']
            rows = [[k['cve'], k['severity'], k['vendor'], str(k['asset_count']), k['kev_name'][:35]]
                   for k in kev_cves[:20]]  # Top 20
            add_styled_table(doc, headers, rows, header_color="B8860B", hyperlink_col=0, font_size=9)

            if len(kev_cves) > 20:
                doc.add_paragraph(f"... and {len(kev_cves) - 20} more KEV CVEs")

        # Assets with Multiple KEV Vulnerabilities - start on own page
        doc.add_page_break()
        doc.add_heading('Assets with Multiple KEV Vulnerabilities', level=2)

        kev_assets = kev_results.get('kev_assets', [])
        # Filter to only assets with 2 or more KEVs
        multi_kev_assets = [a for a in kev_assets if a['kev_count'] >= 2]

        if multi_kev_assets:
            para = doc.add_paragraph()
            para.add_run(f'Total: {len(multi_kev_assets)} assets with multiple KEV vulnerabilities').bold = True
            doc.add_paragraph()

            # Paginate: 40 rows per page (more compact without CVE column)
            ROWS_PER_PAGE = 40
            headers = ['Asset', 'Vendor', 'KEVs', 'Severity']

            for page_start in range(0, len(multi_kev_assets), ROWS_PER_PAGE):
                page_assets = multi_kev_assets[page_start:page_start + ROWS_PER_PAGE]

                rows = [[a['hostname'][:35], a['vendor'], str(a['kev_count']),
                        a['highest_severity']]
                       for a in page_assets]

                add_styled_table(doc, headers, rows, header_color="B8860B", font_size=9)

                # Add page break if more pages to come
                if page_start + ROWS_PER_PAGE < len(multi_kev_assets):
                    doc.add_page_break()
        else:
            doc.add_paragraph('No assets with multiple KEV vulnerabilities found.')

        doc.add_paragraph()

    doc.add_page_break()

    # =========================================================================
    # DEFINITION & ASSET CHANGES
    # =========================================================================
    doc.add_heading('New Vulnerability Definitions', level=2)
    if results['new_definitions']:
        headers = ['Vulnerability', 'CVE', 'Severity', 'Vendor']
        rows = [[d['name'], d['cve'], d['severity'], d['vendor']] for d in results['new_definitions']]
        add_styled_table(doc, headers, rows, header_color="7030A0")
    else:
        doc.add_paragraph('No new vulnerability definitions found in the current dataset.')

    doc.add_paragraph()

    # Removed Vulnerability Definitions
    doc.add_heading('Removed Vulnerability Definitions', level=2)
    if results['removed_definitions']:
        headers = ['Vulnerability', 'CVE', 'Severity', 'Vendor']
        rows = [[d['name'], d['cve'], d['severity'], d['vendor']] for d in results['removed_definitions']]
        add_styled_table(doc, headers, rows, header_color="7030A0")
    else:
        doc.add_paragraph('No vulnerability definitions were removed from the current dataset.')

    doc.add_paragraph()

    # Asset Changes
    doc.add_heading('Asset Inventory Changes', level=2)

    asset_summary = doc.add_paragraph()
    asset_summary.add_run(f"New assets: ").bold = True
    asset_summary.add_run(f"{results['new_asset_count']}")
    asset_summary.add_run(f" | ")
    asset_summary.add_run(f"Removed assets: ").bold = True
    asset_summary.add_run(f"{results['removed_asset_count']}")

    if results['new_assets']:
        doc.add_heading('New Assets', level=3)
        for asset in results['new_assets'][:15]:
            doc.add_paragraph(f"+ {asset}", style='List Bullet')
        if results['new_asset_count'] > 15:
            doc.add_paragraph(f"... and {results['new_asset_count'] - 15} more")

    if results['removed_assets']:
        doc.add_heading('Removed Assets', level=3)
        for asset in results['removed_assets'][:15]:
            doc.add_paragraph(f"- {asset}", style='List Bullet')
        if results['removed_asset_count'] > 15:
            doc.add_paragraph(f"... and {results['removed_asset_count'] - 15} more")

    # =========================================================================
    # CVE CHANGE ANALYSIS (explains week-over-week changes)
    # =========================================================================
    if 'df_prev' in results and 'df_curr' in results:
        doc.add_page_break()
        doc.add_heading('CVE Change Analysis', level=1)

        intro = doc.add_paragraph()
        intro.add_run('This section explains the week-over-week changes by showing which specific devices '
                     'were added or removed for the top-changing CVEs. ')
        intro.add_run('Change values represent the difference in unique device counts.').italic = True

        df_prev = results['df_prev']
        df_curr = results['df_curr']

        for vendor in ['Cisco', 'Palo Alto']:
            doc.add_heading(f'{vendor} - Top Changes', level=2)

            increases, decreases = analyze_top_cve_changes(df_prev, df_curr, vendor=vendor, top_n=3)

            # Top Increases
            if len(increases) > 0:
                inc_para = doc.add_paragraph()
                inc_para.add_run('Top Increases (more devices affected):').bold = True

                for _, row in increases.iterrows():
                    if row['change'] <= 0:
                        continue

                    # Extract CVE ID
                    cve_match = re.search(r'(CVE-\d{4}-\d+)', str(row['cve']))
                    if not cve_match:
                        continue
                    cve_id = cve_match.group(1)

                    # Analyze device changes
                    analysis = analyze_cve_device_changes(df_prev, df_curr, cve_id)

                    # CVE header
                    cve_para = doc.add_paragraph()
                    cve_para.add_run(f"\n{cve_id}").bold = True
                    cve_para.add_run(f" ({row['severity']}) — ")
                    cve_para.add_run(f"Previous: {int(row['prev_devices'])} → Current: {int(row['curr_devices'])} ")
                    change_run = cve_para.add_run(f"(+{int(row['change'])} devices)")
                    change_run.font.color.rgb = RGBColor(192, 0, 0)  # Red

                    # New devices table (limit to 10)
                    if analysis['new_hosts']:
                        new_label = doc.add_paragraph()
                        new_label.add_run(f"New devices ({analysis['new_count']} total):").bold = True

                        display_hosts = analysis['new_hosts'][:10]
                        headers = ['New Devices']
                        rows_data = [[h] for h in display_hosts]
                        add_styled_table(doc, headers, rows_data, header_color="C00000", font_size=8)

                        if analysis['new_count'] > 10:
                            doc.add_paragraph(f"... and {analysis['new_count'] - 10} more new devices")

                doc.add_paragraph()

            # Top Decreases
            if len(decreases) > 0:
                dec_para = doc.add_paragraph()
                dec_para.add_run('Top Decreases (fewer devices affected):').bold = True

                for _, row in decreases.iterrows():
                    if row['change'] >= 0:
                        continue

                    cve_match = re.search(r'(CVE-\d{4}-\d+)', str(row['cve']))
                    if not cve_match:
                        continue
                    cve_id = cve_match.group(1)

                    analysis = analyze_cve_device_changes(df_prev, df_curr, cve_id)

                    cve_para = doc.add_paragraph()
                    cve_para.add_run(f"\n{cve_id}").bold = True
                    cve_para.add_run(f" ({row['severity']}) — ")
                    cve_para.add_run(f"Previous: {int(row['prev_devices'])} → Current: {int(row['curr_devices'])} ")
                    change_run = cve_para.add_run(f"({int(row['change'])} devices)")
                    change_run.font.color.rgb = RGBColor(16, 124, 16)  # Green

                    # Removed devices table (limit to 10)
                    if analysis['removed_hosts']:
                        rem_label = doc.add_paragraph()
                        rem_label.add_run(f"Removed devices ({analysis['removed_count']} total):").bold = True

                        display_hosts = analysis['removed_hosts'][:10]
                        headers = ['Removed Devices']
                        rows_data = [[h] for h in display_hosts]
                        add_styled_table(doc, headers, rows_data, header_color="107C10", font_size=8)

                        if analysis['removed_count'] > 10:
                            doc.add_paragraph(f"... and {analysis['removed_count'] - 10} more removed devices")

            doc.add_paragraph()

    # Save document
    doc.save(output_path)

    # Cleanup temp chart files
    import shutil
    try:
        shutil.rmtree(temp_dir)
    except Exception:
        pass  # Ignore cleanup errors

    print(f"\nReport generated: {output_path}")

def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    prev_file = sys.argv[1]
    curr_file = sys.argv[2]
    output_name = sys.argv[3] if len(sys.argv) > 3 else f"Weekly_Vuln_Report_{datetime.now().strftime('%Y%m%d')}"

    if not output_name.endswith('.docx'):
        output_name += '.docx'

    print(f"Loading previous dataset: {prev_file}")
    df_prev = pd.read_csv(prev_file, encoding='utf-8-sig')

    print(f"Loading current dataset: {curr_file}")
    df_curr = pd.read_csv(curr_file, encoding='utf-8-sig')

    print("Analyzing data...")
    results, df_prev_analyzed, df_curr_analyzed = analyze_data(df_prev, df_curr)

    # Store dataframes for CVE change analysis
    results['df_prev'] = df_prev_analyzed
    results['df_curr'] = df_curr_analyzed

    # Load KEV catalog if available
    print("Checking for KEV catalog...")
    kev_data = load_kev_catalog()
    if kev_data:
        kev_count = len(kev_data.get('vulnerabilities', []))
        print(f"  Found KEV catalog with {kev_count:,} entries")
        print("  Analyzing KEV vulnerabilities...")
        kev_results = analyze_kev_vulnerabilities(df_curr_analyzed, kev_data)
        if kev_results:
            results['kev_analysis'] = kev_results
            print(f"  Found {kev_results['total_kev_cves']} KEV CVEs affecting {kev_results['total_kev_assets']} assets")
        else:
            print("  No KEV vulnerabilities found in dataset")
    else:
        print("  KEV catalog not found (kev_catalog.json)")
        print("  To enable KEV analysis, run: python update_kev_list.py")
        print("  Or download manually from: https://www.cisa.gov/known-exploited-vulnerabilities-catalog")

    print("Generating report...")
    generate_report(results, output_name, os.path.basename(prev_file), os.path.basename(curr_file))

if __name__ == "__main__":
    main()
