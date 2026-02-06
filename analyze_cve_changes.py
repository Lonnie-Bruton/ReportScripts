#!/usr/bin/env python3
"""
CVE Change Analysis Script
==========================
Analyzes the specific device changes for top-changing CVEs between two weeks.
Helps explain why vulnerability counts increased or decreased.

Usage:
    python analyze_cve_changes.py <previous_week.csv> <current_week.csv> [output_name]
"""

import pandas as pd
import sys
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, header_color="2F5496", font_size=9):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        run = header_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(font_size)
        set_cell_shading(header_cells[i], header_color)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
            if row_cells[i].paragraphs[0].runs:
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)

    return table


def normalize_hostname(hostname):
    """Normalize hostname - remove domain suffix."""
    if not hostname or pd.isna(hostname):
        return ""
    clean_hostname = str(hostname).strip()
    ip_pattern = r'^(\d{1,3}\.){3}\d{1,3}$'
    if re.match(ip_pattern, clean_hostname):
        return clean_hostname.lower()
    return clean_hostname.split('.')[0].lower()


def categorize_device_type(hostname, family=None):
    """Categorize device by hostname pattern with family fallback."""
    h = normalize_hostname(hostname)
    f = str(family).lower() if family and pd.notna(family) else ""

    if h:
        # Cisco patterns
        if 'nswan' in h or 'swan' in h or 'nrwan' in h or 'rtr' in h:
            return 'Cisco'
        if re.search(r'ns\d+com01', h):
            return 'Cisco'
        if 'fw' in h and 'nf' not in h:
            return 'Cisco'
        if 'asa' in h:
            return 'Cisco'
        if 'sw' in h and 'wan' not in h:
            return 'Cisco'
        if 'newan' in h:
            return 'Cisco'

        # Palo Alto patterns
        if any(p in h for p in ['nfpan', 'dnfpan', 'nfscada', 'nfdevpan', 'papan',
                                 'nfextpan', 'nfintpan', 'nfcampuspan', 'extpan',
                                 'intpan', 'campuspan', 'devpan']):
            return 'Palo Alto'

        # Other
        if 'nwwbr' in h or 'wwbr' in h:
            return 'Other/General'

    if f:
        if 'cisco' in f:
            return 'Cisco'
        if re.search(r'palo\s*alto', f):
            return 'Palo Alto'

    return 'Other/General'


def analyze_cve_device_changes(df_prev, df_curr, cve_id):
    """Analyze which devices changed for a specific CVE."""

    # Filter to rows containing this CVE
    prev_cve = df_prev[df_prev['definition.cve'].str.contains(cve_id, na=False, regex=False)]
    curr_cve = df_curr[df_curr['definition.cve'].str.contains(cve_id, na=False, regex=False)]

    # Get unique hostnames
    prev_hosts = set(prev_cve['asset.host_name'].dropna().unique())
    curr_hosts = set(curr_cve['asset.host_name'].dropna().unique())

    # Calculate changes
    new_hosts = curr_hosts - prev_hosts
    removed_hosts = prev_hosts - curr_hosts
    unchanged_hosts = prev_hosts & curr_hosts

    return {
        'cve': cve_id,
        'prev_count': len(prev_hosts),
        'curr_count': len(curr_hosts),
        'new_hosts': sorted(new_hosts),
        'removed_hosts': sorted(removed_hosts),
        'unchanged_hosts': sorted(unchanged_hosts),
        'new_count': len(new_hosts),
        'removed_count': len(removed_hosts),
        'unchanged_count': len(unchanged_hosts)
    }


def get_top_changing_cves(df_prev, df_curr, vendor=None, top_n=10):
    """Get the CVEs with the biggest changes."""

    # Add vendor column
    df_prev['vendor'] = df_prev.apply(
        lambda row: categorize_device_type(row['asset.host_name'], row.get('definition.family')),
        axis=1
    )
    df_curr['vendor'] = df_curr.apply(
        lambda row: categorize_device_type(row['asset.host_name'], row.get('definition.family')),
        axis=1
    )

    # Filter by vendor if specified
    if vendor:
        df_prev = df_prev[df_prev['vendor'] == vendor]
        df_curr = df_curr[df_curr['vendor'] == vendor]

    # Count instances per CVE
    prev_counts = df_prev.groupby(['definition.cve', 'definition.name', 'severity']).agg({
        'asset.host_name': 'nunique'
    }).reset_index()
    prev_counts.columns = ['cve', 'name', 'severity', 'prev_devices']

    curr_counts = df_curr.groupby(['definition.cve', 'definition.name', 'severity']).agg({
        'asset.host_name': 'nunique'
    }).reset_index()
    curr_counts.columns = ['cve', 'name', 'severity', 'curr_devices']

    # Merge and calculate change
    merged = prev_counts.merge(curr_counts, on=['cve', 'name', 'severity'], how='outer').fillna(0)
    merged['change'] = merged['curr_devices'] - merged['prev_devices']

    # Get top increases and decreases
    increases = merged.nlargest(top_n, 'change')
    decreases = merged.nsmallest(top_n, 'change')

    return increases, decreases, df_prev, df_curr


def generate_analysis_report(df_prev, df_curr, output_path, prev_name, curr_name):
    """Generate detailed CVE change analysis report."""

    doc = Document()

    # Title
    title = doc.add_heading('CVE Change Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f'Comparing: {prev_name} vs {curr_name}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Explanation
    doc.add_heading('Understanding the Data', level=1)
    explanation = doc.add_paragraph()
    explanation.add_run('What the "Change" numbers mean: ').bold = True
    explanation.add_run('The change value represents the difference in the number of unique devices/assets '
                       'tagged with a specific CVE between the two reporting periods. ')
    explanation.add_run('A large increase typically means one of:\n')
    doc.add_paragraph('New devices were added to scanning scope', style='List Bullet')
    doc.add_paragraph('Existing devices received updated scans that detected previously missed vulnerabilities', style='List Bullet')
    doc.add_paragraph('A new vulnerability definition was added to the scanner', style='List Bullet')
    doc.add_paragraph('Device configurations changed, making them newly vulnerable', style='List Bullet')

    doc.add_page_break()

    # Analyze for Cisco
    for vendor in ['Cisco', 'Palo Alto']:
        doc.add_heading(f'{vendor} - Top CVE Changes', level=1)

        increases, decreases, df_prev_v, df_curr_v = get_top_changing_cves(
            df_prev.copy(), df_curr.copy(), vendor=vendor, top_n=5
        )

        # Top Increases
        doc.add_heading('Top Increases (More Devices Affected)', level=2)

        for _, row in increases.iterrows():
            if row['change'] <= 0:
                continue

            # Extract first CVE ID from the string
            cve_match = re.search(r'(CVE-\d{4}-\d+)', str(row['cve']))
            if not cve_match:
                continue
            cve_id = cve_match.group(1)

            # Analyze this specific CVE
            analysis = analyze_cve_device_changes(df_prev_v, df_curr_v, cve_id)

            # Section header
            header = doc.add_heading(f"{cve_id}", level=3)

            # Summary
            summary = doc.add_paragraph()
            summary.add_run(f"Severity: ").bold = True
            summary.add_run(f"{row['severity']} | ")
            summary.add_run(f"Previous: ").bold = True
            summary.add_run(f"{int(row['prev_devices'])} devices | ")
            summary.add_run(f"Current: ").bold = True
            summary.add_run(f"{int(row['curr_devices'])} devices | ")
            summary.add_run(f"Change: ").bold = True
            summary.add_run(f"+{int(row['change'])} devices")

            # Vulnerability name
            name_para = doc.add_paragraph()
            name_para.add_run("Vulnerability: ").bold = True
            name_para.add_run(str(row['name'])[:100])

            doc.add_paragraph()

            # New devices table
            if analysis['new_hosts']:
                new_para = doc.add_paragraph()
                new_para.add_run(f"NEW devices tagged this week ({analysis['new_count']}):").bold = True

                # Show up to 50 new devices
                display_hosts = analysis['new_hosts'][:50]
                headers = ['#', 'Hostname']
                rows = [[str(i+1), h] for i, h in enumerate(display_hosts)]
                add_styled_table(doc, headers, rows, header_color="C00000", font_size=8)

                if len(analysis['new_hosts']) > 50:
                    doc.add_paragraph(f"... and {len(analysis['new_hosts']) - 50} more new devices")
            else:
                doc.add_paragraph("No new devices (change is due to instance count, not device count)")

            doc.add_paragraph()

            # Removed devices (if any)
            if analysis['removed_hosts']:
                removed_para = doc.add_paragraph()
                removed_para.add_run(f"REMOVED devices (no longer tagged) ({analysis['removed_count']}):").bold = True

                display_hosts = analysis['removed_hosts'][:30]
                headers = ['#', 'Hostname']
                rows = [[str(i+1), h] for i, h in enumerate(display_hosts)]
                add_styled_table(doc, headers, rows, header_color="107C10", font_size=8)

                if len(analysis['removed_hosts']) > 30:
                    doc.add_paragraph(f"... and {len(analysis['removed_hosts']) - 30} more removed devices")

                doc.add_paragraph()

            doc.add_page_break()

        # Top Decreases
        doc.add_heading('Top Decreases (Fewer Devices Affected)', level=2)

        for _, row in decreases.iterrows():
            if row['change'] >= 0:
                continue

            cve_match = re.search(r'(CVE-\d{4}-\d+)', str(row['cve']))
            if not cve_match:
                continue
            cve_id = cve_match.group(1)

            analysis = analyze_cve_device_changes(df_prev_v, df_curr_v, cve_id)

            header = doc.add_heading(f"{cve_id}", level=3)

            summary = doc.add_paragraph()
            summary.add_run(f"Severity: ").bold = True
            summary.add_run(f"{row['severity']} | ")
            summary.add_run(f"Previous: ").bold = True
            summary.add_run(f"{int(row['prev_devices'])} devices | ")
            summary.add_run(f"Current: ").bold = True
            summary.add_run(f"{int(row['curr_devices'])} devices | ")
            summary.add_run(f"Change: ").bold = True
            summary.add_run(f"{int(row['change'])} devices")

            name_para = doc.add_paragraph()
            name_para.add_run("Vulnerability: ").bold = True
            name_para.add_run(str(row['name'])[:100])

            doc.add_paragraph()

            # Removed devices
            if analysis['removed_hosts']:
                removed_para = doc.add_paragraph()
                removed_para.add_run(f"REMOVED devices (no longer vulnerable) ({analysis['removed_count']}):").bold = True

                display_hosts = analysis['removed_hosts'][:50]
                headers = ['#', 'Hostname']
                rows = [[str(i+1), h] for i, h in enumerate(display_hosts)]
                add_styled_table(doc, headers, rows, header_color="107C10", font_size=8)

                if len(analysis['removed_hosts']) > 50:
                    doc.add_paragraph(f"... and {len(analysis['removed_hosts']) - 50} more removed devices")

            doc.add_paragraph()

        doc.add_page_break()

    # Save
    doc.save(output_path)
    print(f"\nAnalysis report generated: {output_path}")


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    prev_file = sys.argv[1]
    curr_file = sys.argv[2]
    output_name = sys.argv[3] if len(sys.argv) > 3 else f"CVE_Change_Analysis_{datetime.now().strftime('%Y%m%d')}"

    if not output_name.endswith('.docx'):
        output_name += '.docx'

    print(f"Loading previous dataset: {prev_file}")
    df_prev = pd.read_csv(prev_file, encoding='utf-8-sig')

    print(f"Loading current dataset: {curr_file}")
    df_curr = pd.read_csv(curr_file, encoding='utf-8-sig')

    print("Analyzing CVE changes...")
    generate_analysis_report(df_prev, df_curr, output_name,
                            os.path.basename(prev_file), os.path.basename(curr_file))


if __name__ == "__main__":
    main()
